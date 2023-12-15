from docx import Document

def replace_text_in_word_file(template_file, new_file, replacements):
    # Open the Word document
    doc = Document(template_file)

    # Function to replace text in a run while keeping the original format
    def replace_text_in_run(run, old_text, new_text):
        if old_text in run.text:
            text = run.text.replace(old_text, new_text)
            run.clear()
            run.add_text(text)

    # Iterate over the first three tables
    for table in doc.tables[:3]:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for old_text, new_text in replacements.items():
                            replace_text_in_run(run, old_text, new_text)

    # Save the new content in a new file
    doc.save(new_file)


if __name__ == '__main__':
    # original_template_file = 'template_edit.docx'
    template_file = 'template_edit.docx'
    replacements = {'{certificado}': 'new_word', '{emisión}': 'new_word', '{póliza}': 'new_word', 
                    '{contratante}': 'SUDAMERIS BANK', '{ruc}': 'new_word', '{nombre}': 'FRANCISCO RUIZ', '{nacimiento}': 'new_word',
                    '{documento}': '4338313', '{nacimiento}': 'new_word', '{domicilio}': 'new_word', 
                    '{amortización}': 'new_word', '{desde}': '11/12/2023', '{hasta}': '23/12/2023', 
                    '{días}': 'new_word', '{fallecimiento}': 'new_word', '{incapacidad}': 'new_word', 
                    '{prima}': 'new_word', '{impuesto}': 'new_word', '{premio}': 'new_word', 
                    '{financiamiento}': 'new_word', '{interés}': 'new_word', '{costo}': 'new_word', 
                    '{final}': 'new_word'}
    
    # Extract the value associated with the key '{nombre}'
    nombre_value = replacements.get('{nombre}', 'default')

    # Form the new file name by concatenating the base name with the value of '{nombre}'
    new_file = f'output/new_file_{nombre_value}.docx'

    replace_text_in_word_file(template_file, new_file, replacements)

    print("Finished")