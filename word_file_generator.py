from docx import Document

def replace_text_in_word_file(template_file, new_file, old_text, new_text):
    # Open the Word document
    doc = Document(template_file)

    # Function to replace text in a run while keeping the original format
    def replace_text_in_run(run, old_text, new_text):
        if old_text in run.text:
            text = run.text.replace(old_text, new_text)
            run.clear()
            run.add_text(text)

    # Iterate over the first three tables
    for table in doc.tables[:3]:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        replace_text_in_run(run, old_text, new_text)

    # Save the new content in a new file
    doc.save(new_file)


if __name__ == '__main__':
    original_template_file = 'template_edit.docx'
    template_file = 'template_edit_test.docx'
    new_file = 'template_edit_test.docx'
    old_text_list = ['{certificado}', '{emisión}','{póliza}', '{contratante}', '{ruc}', '{nacimiento}',
                     '{documento}', '{nacimiento}', '{domicilio}', '{amortización}', '{desde}',
                     '{hasta}', '{días}', '{fallecimiento}', '{incapacidad}', '{prima}',
                     '{impuesto}', '{premio}', '{financiamiento}', '{interés}', '{costo}', '{final}']
    
    for old_text in old_text_list:
        replace_text_in_word_file(template_file, new_file, old_text, 'new_word')

    print("Finished")